"""
ingest.py — Universal Ingestion Pipeline
Handles: URLs, PDFs, YouTube, plain text, folders, .docx
"""

import os
import re
import hashlib
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

import httpx
from bs4 import BeautifulSoup
from markdownify import markdownify as md

from embedder import get_embedder
from utils import chunk_text, get_db, log

_SRC_DIR = os.path.dirname(os.path.abspath(__file__))
_ROOT_DIR = os.path.join(_SRC_DIR, "..")

# ─── Data Model ──────────────────────────────────────────────

@dataclass
class Document:
    content: str
    source: str
    source_type: str            # "url" | "pdf" | "youtube" | "text" | "file"
    title: str = ""
    metadata: dict = field(default_factory=dict)
    doc_id: str = ""

    def __post_init__(self):
        if not self.doc_id:
            self.doc_id = hashlib.md5(
                (self.source + self.content[:100]).encode()
            ).hexdigest()[:12]

# ─── Loaders ─────────────────────────────────────────────────

def load_url(url: str) -> Document:
    """Scrape a webpage and convert to clean Markdown."""
    log(f"Fetching URL: {url}")
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }

    with httpx.Client(follow_redirects=True, timeout=15) as client:
        response = client.get(url, headers=headers)
        response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")

    # Remove noise
    for tag in soup(["script", "style", "nav", "footer", "aside", "iframe"]):
        tag.decompose()

    title = soup.title.string.strip() if soup.title else url
    main = soup.find("article") or soup.find("main") or soup.body
    content = md(str(main), heading_style="ATX") if main else ""
    content = re.sub(r"\n{3,}", "\n\n", content).strip()

    return Document(
        content=content,
        source=url,
        source_type="url",
        title=title,
        metadata={"url": url, "scraped_at": datetime.utcnow().isoformat()},
    )


def load_pdf(path: str) -> Document:
    """Extract text from a PDF file."""
    from PyPDF2 import PdfReader
    log(f"Reading PDF: {path}")

    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages.append(f"[Page {i+1}]\n{text}")

    content = "\n\n".join(pages)
    title = Path(path).stem.replace("_", " ").replace("-", " ").title()

    return Document(
        content=content,
        source=path,
        source_type="pdf",
        title=title,
        metadata={"file": path, "pages": len(reader.pages)},
    )


def load_youtube(url: str) -> Document:
    """Extract transcript from a YouTube video."""
    from youtube_transcript_api import YouTubeTranscriptApi
    from pytube import YouTube
    log(f"Fetching YouTube transcript: {url}")

    video_id = re.search(r"(?:v=|youtu\.be/)([^&\n?#]+)", url)
    if not video_id:
        raise ValueError(f"Could not extract video ID from: {url}")
    vid_id = video_id.group(1)

    # Get title
    try:
        yt = YouTube(url)
        title = yt.title
        channel = yt.author
    except Exception:
        title = f"YouTube: {vid_id}"
        channel = "Unknown"

    # Get transcript
    transcript_list = YouTubeTranscriptApi.get_transcript(vid_id)
    content = " ".join([t["text"] for t in transcript_list])
    content = re.sub(r"\s+", " ", content).strip()

    return Document(
        content=f"# {title}\n\nChannel: {channel}\nSource: {url}\n\n{content}",
        source=url,
        source_type="youtube",
        title=title,
        metadata={"video_id": vid_id, "channel": channel, "url": url},
    )


def load_text(text: str, title: str = "Note") -> Document:
    """Ingest raw text or notes."""
    return Document(
        content=text,
        source=f"text:{hashlib.md5(text.encode()).hexdigest()[:8]}",
        source_type="text",
        title=title,
        metadata={"created_at": datetime.utcnow().isoformat()},
    )


def load_docx(path: str) -> Document:
    """Extract text from a .docx Word file."""
    from docx import Document as DocxDocument
    log(f"Reading DOCX: {path}")
    doc = DocxDocument(path)
    content = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    title = Path(path).stem.replace("_", " ").title()
    return Document(
        content=content,
        source=path,
        source_type="file",
        title=title,
        metadata={"file": path},
    )

# ─── Main Ingest Function ─────────────────────────────────────

def ingest(
    url: Optional[str] = None,
    file: Optional[str] = None,
    youtube: Optional[str] = None,
    text: Optional[str] = None,
    title: str = "",
    folder: Optional[str] = None,
) -> dict:
    """
    Universal entry point. Pass ONE of: url, file, youtube, text, folder.
    Returns a summary dict.
    """
    import chromadb
    from dotenv import load_dotenv
    load_dotenv(os.path.join(_ROOT_DIR, "config", ".env"))

    _chroma_path = os.getenv("CHROMA_PERSIST_DIR", os.path.join(_ROOT_DIR, "data", "chroma_db"))
    if not os.path.isabs(_chroma_path):
        _chroma_path = os.path.join(_ROOT_DIR, _chroma_path)
    chroma = chromadb.PersistentClient(path=_chroma_path)
    collection = chroma.get_or_create_collection(
        name=os.getenv("CHROMA_COLLECTION_NAME", "second_brain"),
        metadata={"hnsw:space": "cosine"},
    )
    embedder = get_embedder()

    # ── Load document ─────────────────────────────────────────
    if folder:
        results = []
        for f in Path(folder).rglob("*"):
            if f.suffix.lower() in [".pdf", ".txt", ".md", ".docx"]:
                result = ingest(file=str(f))
                results.append(result)
        return {"batch": len(results), "results": results}

    if url:
        doc = load_url(url)
    elif youtube:
        doc = load_youtube(youtube)
    elif file:
        ext = Path(file).suffix.lower()
        if ext == ".pdf":
            doc = load_pdf(file)
        elif ext == ".docx":
            doc = load_docx(file)
        elif ext in [".txt", ".md"]:
            doc = load_text(Path(file).read_text(encoding="utf-8", errors="replace"), title=Path(file).stem)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    elif text:
        doc = load_text(text, title=title or "Note")
    else:
        raise ValueError("Provide one of: url, file, youtube, text, folder")

    if title:
        doc.title = title

    # ── Chunk ──────────────────────────────────────────────────
    chunk_size = int(os.getenv("CHUNK_SIZE", 512))
    chunk_overlap = int(os.getenv("CHUNK_OVERLAP", 64))
    chunks = chunk_text(doc.content, chunk_size=chunk_size, overlap=chunk_overlap)

    if not chunks:
        return {"status": "empty", "doc_id": doc.doc_id}

    # ── Embed & Store ──────────────────────────────────────────
    embeddings = embedder.embed(chunks)

    ids = [f"{doc.doc_id}_{i}" for i in range(len(chunks))]
    metadatas = [
        {
            **doc.metadata,
            "doc_id": doc.doc_id,
            "title": doc.title,
            "source": doc.source,
            "source_type": doc.source_type,
            "chunk_index": i,
            "total_chunks": len(chunks),
        }
        for i in range(len(chunks))
    ]

    collection.add(
        ids=ids,
        documents=chunks,
        embeddings=embeddings.tolist(),
        metadatas=metadatas,
    )

    # Write to SQL documents table
    try:
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT OR REPLACE INTO documents (doc_id, title, source, source_type, chunks, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (doc.doc_id, doc.title, doc.source, doc.source_type, len(chunks), datetime.utcnow().isoformat()))
        conn.commit()
        conn.close()
    except Exception as db_err:
        log(f"[Ingest DB Error] {db_err}", "WARN")

    log(f"Ingested: '{doc.title}' -> {len(chunks)} chunks stored.")
    return {
        "status": "ok",
        "doc_id": doc.doc_id,
        "title": doc.title,
        "chunks": len(chunks),
        "source_type": doc.source_type,
    }
